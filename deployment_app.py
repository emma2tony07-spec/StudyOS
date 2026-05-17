# ═══════════════════════════════════════════════════════════════
#  StudyOS Backend — Render.com Edition
#
#  Setup on Render:
#  1. Push this file to a GitHub repo
#  2. Create new Web Service on render.com
#  3. Build command:  pip install -r requirements.txt
#  4. Start command:  gunicorn app:app --workers 1 --threads 2 --timeout 120 --bind 0.0.0.0:$PORT
#  5. Set environment variables in Render dashboard:
#       MASTER_OPENROUTER_KEY = sk-or-v1-...
#       GMAIL_ADDRESS         = your@gmail.com
#       GMAIL_APP_PASS        = xxxxxxxxxxxxxxxx  (no spaces!)
#
#  Note: Render free tier spins down after inactivity.
#  The app will wake on first request (cold start ~30s).
# ═══════════════════════════════════════════════════════════════

from flask import Flask, request, jsonify
from flask_cors import CORS
import json, os, re, smtplib, threading, time, requests
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from datetime import datetime, date
from werkzeug.utils import secure_filename
from collections import Counter

app = Flask(__name__)

# On Render, allow your deployed frontend URL
# Replace with your actual frontend URL once deployed
ALLOWED_ORIGINS = [
    "http://localhost:8080",
    "http://localhost:5500",
    "https://aretistudyos.netlify.app/",
    "https://emma2tony07-spec.github.io/StudyOS/",
    "https://your-studyos-frontend.vercel.app",   # UPDATE THIS
]

CORS(app, origins=ALLOWED_ORIGINS, supports_credentials=True)

# ═══════════════════════════════════════════════════════════════
#  CONFIGURATION — pulled from environment variables on Render
# ═══════════════════════════════════════════════════════════════
GMAIL_ADDRESS         = os.environ.get("GMAIL_ADDRESS", "").strip()
GMAIL_APP_PASS        = os.environ.get("GMAIL_APP_PASS", "").replace(" ", "").strip()
FROM_NAME             = "StudyOS"
MASTER_OPENROUTER_KEY = os.environ.get("MASTER_OPENROUTER_KEY", "").strip()
OPENROUTER_BASE_URL   = "https://openrouter.ai/api/v1"
DAILY_FREE_LIMIT      = int(os.environ.get("DAILY_FREE_LIMIT", "20"))

AVAILABLE_MODELS = [
    "poolside/laguna-xs.2:free",
    "poolside/laguna-m.1:free",
    "nvidia/nemotron-3-nano-30b-a3b:free",
    "nvidia/nemotron-nano-9b-v2:free",
    "openai/gpt-oss-120b:free",
    "openai/gpt-oss-20b:free",
    "z-ai/glm-4.5-air:free",
    "google/gemma-4-26b-a4b-it:free",
    "google/gemma-4-31b-it:free",
    "minimax/minimax-m2.5:free",
    "qwen/qwen3-next-80b-a3b-instruct:free",
    "meta-llama/llama-3.3-70b-instruct:free",
    "meta-llama/llama-3.2-3b-instruct:free",
]

DEFAULT_CONFIG = {
    "analyze":   "openai/gpt-oss-120b:free",
    "summarize": "google/gemma-4-26b-a4b-it:free",
    "grade":     "nvidia/nemotron-3-nano-30b-a3b:free",
    "simplify":  "google/gemma-4-26b-a4b-it:free",
    "visualize": "z-ai/glm-4.5-air:free",
}

# On Render, use /tmp for ephemeral storage
# (filesystem resets on redeploy — rate limits and config reset too)
TMP_DIR         = "/tmp/studyos"
os.makedirs(TMP_DIR, exist_ok=True)
CONFIG_FILE     = os.path.join(TMP_DIR, "studyos_config.json")
RATE_LIMIT_FILE = os.path.join(TMP_DIR, "rate_limits.json")
UPLOAD_FOLDER   = os.path.join(TMP_DIR, "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ═══════════════════════════════════════════════════════════════
#  STARTUP VALIDATION
# ═══════════════════════════════════════════════════════════════
print("=" * 60)
print("🚀 StudyOS Backend Starting...")
print("=" * 60)

if not MASTER_OPENROUTER_KEY:
    print("⚠️  WARNING: MASTER_OPENROUTER_KEY is not set!")
    print("   Users will need to provide their own API keys.")
else:
    print(f"✅ Master OpenRouter key configured (starts with: {MASTER_OPENROUTER_KEY[:15]}...)")

if not GMAIL_ADDRESS or not GMAIL_APP_PASS:
    print("⚠️  WARNING: Email notifications not configured")
    print("   Set GMAIL_ADDRESS and GMAIL_APP_PASS in Render environment variables.")
else:
    print(f"✅ Email configured for: {GMAIL_ADDRESS}")

print(f"📊 Daily free limit: {DAILY_FREE_LIMIT} requests per user")
print(f"💾 Storage directory: {TMP_DIR}")
print("=" * 60)

# ═══════════════════════════════════════════════════════════════
#  CONFIG PERSISTENCE
# ═══════════════════════════════════════════════════════════════
config_lock = threading.Lock()

def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r") as f:
                saved  = json.load(f)
                config = DEFAULT_CONFIG.copy()
                config.update(saved)
                return config
        except Exception as e:
            print(f"⚠️  Error loading config: {e}")
    return DEFAULT_CONFIG.copy()

def save_config(config):
    try:
        with open(CONFIG_FILE, "w") as f:
            json.dump(config, f, indent=2)
    except Exception as e:
        print(f"⚠️  Error saving config: {e}")

model_config = load_config()

# ═══════════════════════════════════════════════════════════════
#  RATE LIMITING
# ═══════════════════════════════════════════════════════════════
rate_lock = threading.Lock()

def load_rate_limits():
    if os.path.exists(RATE_LIMIT_FILE):
        try:
            with open(RATE_LIMIT_FILE, "r") as f:
                return json.load(f)
        except Exception as e:
            print(f"⚠️  Error loading rate limits: {e}")
    return {}

def save_rate_limits(data):
    try:
        with open(RATE_LIMIT_FILE, "w") as f:
            json.dump(data, f, indent=2)
    except Exception as e:
        print(f"⚠️  Error saving rate limits: {e}")

def get_user_usage(uid):
    today = str(date.today())
    with rate_lock:
        limits = load_rate_limits()
        user   = limits.get(uid, {})
        if user.get("date") != today:
            return 0, False
        count = user.get("count", 0)
        return count, count >= DAILY_FREE_LIMIT

def increment_usage(uid):
    today = str(date.today())
    with rate_lock:
        limits = load_rate_limits()
        user   = limits.get(uid, {})
        if user.get("date") != today:
            user = {"date": today, "count": 0}
        user["count"]  = user.get("count", 0) + 1
        limits[uid]    = user
        save_rate_limits(limits)
        return user["count"]

def resolve_api_key(request_obj):
    personal_key = request_obj.headers.get("X-User-Api-Key", "").strip()
    if personal_key:
        return personal_key, True
    return MASTER_OPENROUTER_KEY, False

def check_and_charge(request_obj):
    uid = request_obj.headers.get("X-User-Uid", "").strip()
    if not uid:
        return None, None, False, (jsonify({"error": "Missing user ID. Please sign in again."}), 401)
    api_key, is_personal = resolve_api_key(request_obj)
    if is_personal:
        return uid, api_key, True, None
    if not api_key:
        return uid, None, False, (jsonify({
            "error": "no_api_key",
            "message": "No API key configured. Please add your OpenRouter API key in AI Model Settings."
        }), 401)
    count, limited = get_user_usage(uid)
    if limited:
        return uid, api_key, False, (jsonify({
            "error":   "daily_limit_reached",
            "message": f"You've used all {DAILY_FREE_LIMIT} free daily requests. Add your own OpenRouter API key for unlimited access.",
            "usage":   count,
            "limit":   DAILY_FREE_LIMIT,
            "resets":  "midnight"
        }), 429)
    increment_usage(uid)
    return uid, api_key, False, None

# ═══════════════════════════════════════════════════════════════
#  FILE UPLOADS
# ═══════════════════════════════════════════════════════════════
ALLOWED_EXTENSIONS = {'pdf', 'txt', 'docx'}
app.config['UPLOAD_FOLDER']      = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024  # 20MB max

# ═══════════════════════════════════════════════════════════════
#  PDF / DOCUMENT EXTRACTION
# ═══════════════════════════════════════════════════════════════
def is_scanned_pdf(filepath):
    try:
        import fitz
        doc        = fitz.open(filepath)
        pages      = min(5, len(doc))
        total_text = 0
        for i in range(pages):
            total_text += len(doc[i].get_text().strip())
        doc.close()
        return (total_text / max(pages, 1)) < 50
    except Exception as e:
        print(f"⚠️  Error checking if PDF is scanned: {e}")
        return False

def extract_repeating_lines(pages_text, threshold=0.4):
    if len(pages_text) < 3:
        return set()
    line_counts = Counter()
    for page in pages_text:
        seen_on_page = set()
        for line in page.splitlines():
            stripped = line.strip()
            if len(stripped) > 5:
                norm = re.sub(r'\s+', ' ', stripped.lower())
                if norm not in seen_on_page:
                    line_counts[norm] += 1
                    seen_on_page.add(norm)
    total_pages = len(pages_text)
    repeating   = set()
    for line, count in line_counts.items():
        if count / total_pages >= threshold:
            repeating.add(line)
    return repeating

def clean_page_text(page_text, repeating_lines):
    lines   = page_text.splitlines()
    cleaned = []
    for line in lines:
        stripped = line.strip()
        norm     = re.sub(r'\s+', ' ', stripped.lower())
        if norm in repeating_lines:
            continue
        if re.match(r'^\d{1,3}$', stripped):
            continue
        cleaned.append(stripped)
    result     = []
    prev_blank = False
    for line in cleaned:
        if not line:
            if not prev_blank:
                result.append('')
            prev_blank = True
        else:
            result.append(line)
            prev_blank = False
    return '\n'.join(result).strip()

def extract_text_from_pdf(filepath):
    try:
        import fitz
    except ImportError:
        return None, "pymupdf not installed. Run: pip install pymupdf"

    if is_scanned_pdf(filepath):
        return None, (
            "This PDF appears to be a scanned image and contains no extractable text. "
            "Please convert it to a digital PDF first using an OCR tool such as:\n"
            "• Adobe Acrobat (OCR option)\n"
            "• Microsoft Lens (free phone app)\n"
            "• Google Drive: upload PDF → open with Google Docs → download as PDF\n"
            "• smallpdf.com → OCR PDF tool\n"
            "Then re-upload the converted file."
        )

    try:
        doc        = fitz.open(filepath)
        pages_text = []
        for page in doc:
            pages_text.append(page.get_text("text"))
        doc.close()

        repeating     = extract_repeating_lines(pages_text)
        cleaned_pages = []
        for page_text in pages_text:
            cleaned = clean_page_text(page_text, repeating)
            if cleaned:
                cleaned_pages.append(cleaned)

        full_text = '\n\n'.join(cleaned_pages)

        if len(full_text.strip()) < 100:
            return None, (
                "Very little text could be extracted from this PDF. "
                "If it contains scanned images or is password-protected, "
                "please convert it to a digital PDF first."
            )
        return full_text, None
    except Exception as e:
        print(f"❌ PDF extraction error: {e}")
        return None, f"PDF extraction failed: {str(e)}"

def extract_text_from_txt(filepath):
    for enc in ('utf-8', 'latin-1', 'cp1252'):
        try:
            with open(filepath, 'r', encoding=enc) as f:
                return f.read(), None
        except UnicodeDecodeError:
            continue
    return None, "Could not decode text file. Try saving it as UTF-8."

def extract_text_from_docx(filepath):
    try:
        import docx
        doc   = docx.Document(filepath)
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(' | '.join(cells))
        text = '\n'.join(lines)
        if not text.strip():
            return None, "The DOCX file appears to be empty."
        return text, None
    except Exception as e:
        print(f"❌ DOCX extraction error: {e}")
        return None, f"DOCX extraction failed: {str(e)}"

def get_text(filepath):
    ext = filepath.rsplit('.', 1)[1].lower()
    if ext == 'pdf':  return extract_text_from_pdf(filepath)
    if ext == 'txt':  return extract_text_from_txt(filepath)
    if ext == 'docx': return extract_text_from_docx(filepath)
    return None, "Unsupported file type."

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# ═══════════════════════════════════════════════════════════════
#  OPENROUTER CALLER
# ═══════════════════════════════════════════════════════════════
def call_openrouter(messages, feature="summarize", api_key=None, expect_json=True):
    model = model_config.get(feature, DEFAULT_CONFIG[feature])
    key   = api_key or MASTER_OPENROUTER_KEY

    if not key:
        raise Exception("No OpenRouter API key configured.")

    headers = {
        "Authorization": f"Bearer {key}",
        "Content-Type":  "application/json",
        "HTTP-Referer":  "https://studyos.app",
        "X-Title":       "StudyOS",
    }
    payload = {
        "model":       model,
        "messages":    messages,
        "temperature": 0.7,
    }
    if expect_json:
        payload["response_format"] = {"type": "json_object"}

    try:
        response = requests.post(
            f"{OPENROUTER_BASE_URL}/chat/completions",
            headers=headers, json=payload, timeout=120
        )
        if response.status_code == 429:
            raise Exception("Rate limit hit on OpenRouter. Please try again in a moment.")
        if response.status_code == 401:
            raise Exception("Invalid OpenRouter API key. Please check your key in AI Model Settings.")
        if response.status_code == 402:
            raise Exception("OpenRouter account has insufficient credits.")
        response.raise_for_status()
        data    = response.json()
        content = data["choices"][0]["message"]["content"]
        if expect_json:
            clean = re.sub(r'^```(?:json)?\s*|\s*```$', '', content.strip())
            try:
                return json.loads(clean)
            except json.JSONDecodeError:
                match = re.search(r'\{.*\}', clean, re.DOTALL)
                if match:
                    return json.loads(match.group())
                raise Exception("AI returned malformed JSON. Try again.")
        return content
    except requests.exceptions.Timeout:
        raise Exception("OpenRouter request timed out. Try again.")
    except requests.exceptions.ConnectionError:
        raise Exception("Cannot connect to OpenRouter. Check your internet connection.")
    except Exception as e:
        if "No OpenRouter API key" in str(e):
            raise
        print(f"❌ OpenRouter error: {e}")
        raise

# ═══════════════════════════════════════════════════════════════
#  REMINDER STATE + EMAIL
# ═══════════════════════════════════════════════════════════════
pending_reminders = {}
reminder_lock     = threading.Lock()

TYPE_LABELS = {
    "reading":    {"label": "Reading Session",  "emoji": "📖"},
    "quiz":       {"label": "Quiz",             "emoji": "✏️"},
    "assignment": {"label": "Assignment",       "emoji": "📝"},
    "revision":   {"label": "Revision Session", "emoji": "🔄"},
}

def build_email_html(to_name, task, course_name, subject_line, intro_line):
    cfg      = TYPE_LABELS.get(task.get("type",""), {"label": task.get("type","Task"), "emoji": "📚"})
    task_dt  = datetime.fromisoformat(f"{task['date']}T{task['time']}")
    date_str = task_dt.strftime("%A, %B %-d, %Y")
    time_str = task_dt.strftime("%-I:%M %p")
    duration = int(task.get("duration", 30))
    dur_str  = f"{duration // 60} hour{'s' if duration > 60 else ''}" if duration >= 60 else f"{duration} minutes"
    course_row = f'<p style="margin:0 0 8px;color:#A09A92;font-size:14px"><strong style="color:#F0EDE8">📚 Course:</strong> {course_name}</p>' if course_name else ""
    notes_row  = f'<p style="margin:0 0 8px;color:#A09A92;font-size:14px"><strong style="color:#F0EDE8">💬 Notes:</strong> {task["notes"]}</p>' if task.get("notes") else ""
    return f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8"/></head>
<body style="margin:0;padding:0;background:#0D0D0D;font-family:'Helvetica Neue',Helvetica,Arial,sans-serif">
  <table width="100%" cellpadding="0" cellspacing="0" style="background:#0D0D0D;padding:40px 20px">
    <tr><td align="center">
      <table width="560" cellpadding="0" cellspacing="0"
             style="background:#141414;border-radius:16px;overflow:hidden;border:1px solid rgba(255,255,255,0.08)">
        <tr><td style="background:#9B1D20;padding:28px 36px">
            <p style="margin:0 0 10px;display:inline-block;background:rgba(255,255,255,0.18);border-radius:8px;padding:4px 12px;font-size:11px;color:white;font-weight:700;letter-spacing:0.08em;text-transform:uppercase">{subject_line}</p>
            <h1 style="margin:0;color:white;font-size:24px;font-weight:700">{cfg['emoji']} {cfg['label']} Time!</h1>
        </td></tr>
        <tr><td style="padding:32px 36px">
            <p style="margin:0 0 24px;font-size:15px;color:#A09A92;line-height:1.6">
              Hi <strong style="color:#F0EDE8">{to_name}</strong>, {intro_line}
            </p>
            <div style="background:#1C1C1C;border-radius:12px;padding:20px 24px;border-left:4px solid #9B1D20;margin-bottom:28px">
              <h2 style="margin:0 0 14px;color:#F0EDE8;font-size:17px;font-weight:600">{task['title']}</h2>
              <p style="margin:0 0 8px;color:#A09A92;font-size:14px"><strong style="color:#F0EDE8">📅 When:</strong> {date_str} at {time_str}</p>
              <p style="margin:0 0 8px;color:#A09A92;font-size:14px"><strong style="color:#F0EDE8">⏱ Duration:</strong> {dur_str}</p>
              {course_row}{notes_row}
            </div>
            <div style="text-align:center">
              <a href="https://your-studyos-frontend.netlify.app"
                 style="display:inline-block;background:#9B1D20;color:white;text-decoration:none;padding:13px 32px;border-radius:10px;font-size:15px;font-weight:600">
                Open StudyOS →
              </a>
            </div>
        </td></tr>
        <tr><td style="padding:18px 36px;border-top:1px solid rgba(255,255,255,0.06)">
            <p style="margin:0;font-size:12px;color:#6B6560;text-align:center">StudyOS · Sent because you scheduled a study task.</p>
        </td></tr>
      </table>
    </td></tr>
  </table>
</body></html>"""

def build_email_text(to_name, task, course_name, intro_line):
    cfg      = TYPE_LABELS.get(task.get("type",""), {"label": task.get("type","Task"), "emoji": "📚"})
    task_dt  = datetime.fromisoformat(f"{task['date']}T{task['time']}")
    date_str = task_dt.strftime("%A, %B %-d, %Y")
    time_str = task_dt.strftime("%-I:%M %p")
    duration = int(task.get("duration", 30))
    dur_str  = f"{duration // 60}h" if duration >= 60 else f"{duration}m"
    lines    = [f"Hi {to_name}, {intro_line}", "",
                f"{cfg['emoji']} {cfg['label']}: {task['title']}",
                f"📅 {date_str} at {time_str} ({dur_str})"]
    if course_name:       lines.append(f"📚 Course: {course_name}")
    if task.get("notes"): lines.append(f"💬 Notes: {task['notes']}")
    lines   += ["", "Open StudyOS to get started.", "", "— StudyOS"]
    return "\n".join(lines)

def send_email(to_address, to_name, subject, task, course_name, subject_line_badge, intro_line):
    if not GMAIL_ADDRESS or not GMAIL_APP_PASS:
        print("⚠️  Gmail not configured — skipping")
        return False
    try:
        msg            = MIMEMultipart("alternative")
        msg["Subject"] = subject
        msg["From"]    = f"{FROM_NAME} <{GMAIL_ADDRESS}>"
        msg["To"]      = to_address
        msg.attach(MIMEText(build_email_text(to_name, task, course_name, intro_line), "plain"))
        msg.attach(MIMEText(build_email_html(to_name, task, course_name, subject_line_badge, intro_line), "html"))
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(GMAIL_ADDRESS, GMAIL_APP_PASS)
            server.sendmail(GMAIL_ADDRESS, to_address, msg.as_string())
        print(f"📧 Email sent → {to_address}")
        return True
    except smtplib.SMTPAuthenticationError:
        print("❌ Gmail auth failed. Check your GMAIL_ADDRESS and GMAIL_APP_PASS.")
        return False
    except Exception as e:
        print(f"❌ Email error: {e}")
        return False

def reminder_checker():
    print("⏰ Reminder thread started")
    while True:
        time.sleep(60)
        now = datetime.now()
        with reminder_lock:
            for task_id, entry in list(pending_reminders.items()):
                if entry["reminder_sent"]: continue
                task = entry["task"]
                if not task.get("date") or not task.get("time"): continue
                try:
                    task_dt  = datetime.fromisoformat(f"{task['date']}T{task['time']}")
                    diff_min = (task_dt - now).total_seconds() / 60
                    if 0 <= diff_min <= 5:
                        send_email(
                            to_address=entry["user_email"], to_name=entry["user_name"],
                            subject=f"⏰ Starting soon: {task['title']} at {task_dt.strftime('%-I:%M %p')}",
                            task=task, course_name=entry.get("course_name"),
                            subject_line_badge="5-Minute Reminder",
                            intro_line=f"your <strong style='color:#F0EDE8'>{task['title']}</strong> session starts in about 5 minutes."
                        )
                        entry["reminder_sent"] = True
                    elif diff_min < 0:
                        entry["reminder_sent"] = True
                except Exception as e:
                    print(f"Reminder error {task_id}: {e}")

threading.Thread(target=reminder_checker, daemon=True).start()

# ═══════════════════════════════════════════════════════════════
#  AI FUNCTIONS
# ═══════════════════════════════════════════════════════════════
def ai_detect_chapters(text, api_key):
    if len(text) > 18000:
        cut       = text[:18000]
        last_para = cut.rfind('\n\n')
        text      = cut[:last_para] if last_para > 10000 else cut
    messages = [
        {"role": "system", "content": (
            "You are an expert document analyser. "
            "Read the document and detect logical chapter or section boundaries. "
            "Each chapter should contain the full relevant text for that section. "
            "If the document has no clear chapters, split it into logical topic sections. "
            'Return ONLY valid JSON: {"chapters": [{"title": "string", "content": "string"}]}'
        )},
        {"role": "user", "content": text}
    ]
    return call_openrouter(messages, feature="analyze", api_key=api_key, expect_json=True)

def ai_summarize_and_question(text, api_key):
    if len(text) > 15000: text = text[:15000]
    messages = [
        {"role": "system", "content": (
            "You are an expert tutor. Given study material:\n"
            "1. Write a concise summary (3-5 sentences).\n"
            "2. Generate exactly 3 practice questions.\n"
            'Return ONLY valid JSON: {"summary": "string", "questions": ["string","string","string"]}'
        )},
        {"role": "user", "content": text}
    ]
    return call_openrouter(messages, feature="summarize", api_key=api_key, expect_json=True)

def ai_grade_answer(chapter_text, student_answer, api_key):
    if len(chapter_text) > 12000: chapter_text = chapter_text[:12000]
    messages = [
        {"role": "system", "content": (
            "You are a strict but fair academic evaluator. "
            'Return ONLY valid JSON: {"score": <integer 0-10>, "feedback": "string"} '
            "Score: 0-5=poor, 6-7=adequate, 8-9=good, 10=excellent."
        )},
        {"role": "user", "content": f"Chapter:\n{chapter_text}\n\nStudent Answer:\n{student_answer}"}
    ]
    return call_openrouter(messages, feature="grade", api_key=api_key, expect_json=True)

def ai_simplify(text, api_key):
    if len(text) > 15000: text = text[:15000]
    messages = [
        {"role": "system", "content": (
            "You are a brilliant teacher. Rewrite content simply for a 15-year-old. "
            "Use at least 2 real-world analogies and 1 concrete example. "
            'Return ONLY valid JSON: {"title": "string", "sections": [{"heading": "string", "body": "string", "analogy": "string or null", "example": "string or null"}]}'
        )},
        {"role": "user", "content": text}
    ]
    return call_openrouter(messages, feature="simplify", api_key=api_key, expect_json=True)

def ai_visualize(text, api_key):
    if len(text) > 15000: text = text[:15000]
    messages = [
        {"role": "system", "content": (
            "Extract structured knowledge. Return ONLY valid JSON:\n"
            '{"topic":"string","key_concepts":[{"term":"string","definition":"string","importance":"high|medium|low"}],'
            '"relationships":[{"from":"string","to":"string","label":"string"}],'
            '"stats":[{"label":"string","value":"string or number","unit":"string or null"}],'
            '"timeline":[{"step":number,"event":"string","detail":"string"}],'
            '"flashcards":[{"front":"string","back":"string"}]}'
        )},
        {"role": "user", "content": text}
    ]
    return call_openrouter(messages, feature="visualize", api_key=api_key, expect_json=True)

# ═══════════════════════════════════════════════════════════════
#  ROUTES
# ═══════════════════════════════════════════════════════════════
@app.route('/health', methods=['GET'])
def health():
    return jsonify({
        "status": "healthy",
        "provider": "openrouter",
        "model_config": model_config,
        "daily_limit": DAILY_FREE_LIMIT,
        "master_key_set": bool(MASTER_OPENROUTER_KEY),
        "email_configured": bool(GMAIL_ADDRESS and GMAIL_APP_PASS),
        "timestamp": datetime.now().isoformat()
    })

@app.route('/usage', methods=['GET'])
def get_usage():
    uid = request.headers.get("X-User-Uid", "").strip()
    if not uid:
        return jsonify({"error": "Missing user ID"}), 401
    count, limited = get_user_usage(uid)
    personal_key   = request.headers.get("X-User-Api-Key", "").strip()
    return jsonify({
        "uid": uid,
        "used": count,
        "limit": DAILY_FREE_LIMIT,
        "remaining": max(0, DAILY_FREE_LIMIT - count),
        "limited": limited,
        "has_own_key": bool(personal_key),
        "resets": "midnight"
    })

@app.route('/models', methods=['GET'])
def get_models():
    return jsonify({
        "available": AVAILABLE_MODELS,
        "config": model_config,
        "features": ["analyze", "summarize", "grade", "simplify", "visualize"]
    })

@app.route('/models/config', methods=['POST'])
def set_model_config():
    global model_config
    data = request.get_json()
    if not data:
        return jsonify({"error": "No data provided"}), 400
    
    with config_lock:
        for feature in ["analyze", "summarize", "grade", "simplify", "visualize"]:
            if feature in data and data[feature] in AVAILABLE_MODELS:
                model_config[feature] = data[feature]
        save_config(model_config)
    
    return jsonify({"ok": True, "config": model_config})

@app.route('/validate-key', methods=['POST'])
def validate_key():
    data = request.get_json()
    key  = (data or {}).get("key", "").strip()
    if not key:
        return jsonify({"valid": False, "error": "No key provided"}), 400
    
    try:
        headers = {
            "Authorization": f"Bearer {key}",
            "Content-Type": "application/json"
        }
        resp = requests.get(
            f"{OPENROUTER_BASE_URL}/models",
            headers=headers,
            timeout=10
        )
        if resp.status_code == 200:
            return jsonify({"valid": True})
        elif resp.status_code == 401:
            return jsonify({"valid": False, "error": "Invalid API key."})
        else:
            return jsonify({
                "valid": False,
                "error": f"OpenRouter returned status {resp.status_code}."
            })
    except requests.exceptions.Timeout:
        return jsonify({"valid": False, "error": "Connection timed out. Check your internet."}), 500
    except requests.exceptions.ConnectionError:
        return jsonify({"valid": False, "error": "Cannot connect to OpenRouter. Check your internet."}), 500
    except Exception as e:
        return jsonify({"valid": False, "error": str(e)}), 500

@app.route('/analyze', methods=['POST'])
def analyze_file():
    uid, api_key, is_personal, err = check_and_charge(request)
    if err:
        return err
    
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files['file']
        if not file.filename or not allowed_file(file.filename):
            return jsonify({
                "error": "Invalid file type. Please upload PDF, TXT, or DOCX files only."
            }), 400
        
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        text, extract_error = get_text(filepath)
        
        # Clean up uploaded file
        try:
            os.remove(filepath)
        except Exception as e:
            print(f"⚠️  Could not remove temp file: {e}")
        
        if extract_error:
            return jsonify({"error": extract_error}), 400
        
        if not text or len(text.strip()) < 50:
            return jsonify({
                "error": "Could not extract enough text from the file. The file might be empty or corrupted."
            }), 400        
        print(f"📄 Analyzing {len(text)} chars | model: {model_config['analyze']} | uid: {uid}")
        result = ai_detect_chapters(text, api_key)
        chapters = result.get("chapters", [])
        
        if not chapters:
            raise ValueError("AI returned no chapters. Please try again.")
        
        return jsonify({
            "chapters": chapters,
            "total_chapters": len(chapters)
        })
    except Exception as e:
        print(f"❌ /analyze error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/summarize', methods=['POST'])
def summarize():
    uid, api_key, is_personal, err = check_and_charge(request)
    if err:
        return err
    
    try:
        file = request.files.get('file')
        if not file:
            return jsonify({"error": "No text content provided"}), 400
        
        text = file.read().decode('utf-8', errors='ignore')
        if not text.strip():
            return jsonify({"error": "Empty chapter content"}), 400
        
        print(f"📝 Summarizing {len(text)} chars | uid: {uid}")
        result = ai_summarize_and_question(text, api_key)
        return jsonify({
            "summary": result.get("summary", "No summary generated."),
            "questions": result.get("questions", [])
        })
    except Exception as e:
        print(f"❌ /summarize error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/grade', methods=['POST'])
def grade():
    uid, api_key, is_personal, err = check_and_charge(request)
    if err:
        return err
    
    try:
        chapter_text = request.form.get('chapter_text', '').strip()
        student_answer = request.form.get('answer', '').strip()
        
        if not student_answer:
            return jsonify({"error": "No answer provided"}), 400
        
        print(f"🎯 Grading answer | uid: {uid}")
        result = ai_grade_answer(chapter_text, student_answer, api_key)
        score = max(0, min(10, int(result.get("score", 5))))
        feedback = result.get("feedback", "Good effort!")
        
        return jsonify({
            "score": score,
            "feedback": feedback
        })
    except Exception as e:
        print(f"❌ /grade error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/simplify', methods=['POST'])
def simplify():
    uid, api_key, is_personal, err = check_and_charge(request)
    if err:
        return err
    
    try:
        file = request.files.get('file')
        if not file:
            return jsonify({"error": "No text content provided"}), 400
        
        text = file.read().decode('utf-8', errors='ignore')
        if not text.strip():
            return jsonify({"error": "Empty chapter content"}), 400
        
        print(f"📚 Simplifying {len(text)} chars | uid: {uid}")
        result = ai_simplify(text, api_key)
        return jsonify(result)
    except Exception as e:
        print(f"❌ /simplify error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/visualize-data', methods=['POST'])
def visualize_data():
    uid, api_key, is_personal, err = check_and_charge(request)
    if err:
        return err
    
    try:
        file = request.files.get('file')
        if not file:
            return jsonify({"error": "No text content provided"}), 400
        
        text = file.read().decode('utf-8', errors='ignore')
        if not text.strip():
            return jsonify({"error": "Empty chapter content"}), 400
        
        print(f"🧠 Visualizing {len(text)} chars | uid: {uid}")
        result = ai_visualize(text, api_key)
        return jsonify(result)
    except Exception as e:
        print(f"❌ /visualize-data error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/send-confirmation', methods=['POST'])
def send_confirmation():
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400
        
        task = data.get("task")
        user_email = data.get("user_email", "").strip()
        user_name = data.get("user_name", "Student")
        course_name = data.get("course_name", None)
        
        if not task or not user_email:
            return jsonify({"error": "Missing task or user_email"}), 400
        
        if not task.get("date") or not task.get("time"):
            return jsonify({
                "ok": True,
                "message": "No time set — confirmation email skipped"
            }), 200
        
        task_dt = datetime.fromisoformat(f"{task['date']}T{task['time']}")
        time_str = task_dt.strftime("%-I:%M %p")
        
        email_sent = send_email(
            to_address=user_email,
            to_name=user_name,
            subject=f"✅ Task Scheduled: {task['title']} at {time_str}",
            task=task,
            course_name=course_name,
            subject_line_badge="Task Confirmed",
            intro_line=f"your task <strong style='color:#F0EDE8'>{task['title']}</strong> has been scheduled."
        )
        
        # Store for reminder even if email failed
        task_id = task.get("id", str(time.time()))
        with reminder_lock:
            pending_reminders[task_id] = {
                "task": task,
                "user_email": user_email,
                "user_name": user_name,
                "course_name": course_name,
                "reminder_sent": False
            }
        
        return jsonify({
            "ok": True,
            "email_sent": email_sent
        })
    except Exception as e:
        print(f"❌ /send-confirmation error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/cancel-reminder', methods=['POST'])
def cancel_reminder():
    try:
        data = request.get_json()
        task_id = data.get("task_id") if data else None
        
        if task_id:
            with reminder_lock:
                removed = pending_reminders.pop(task_id, None)
                if removed:
                    print(f"🔕 Reminder cancelled for task: {task_id}")
        
        return jsonify({"ok": True})
    except Exception as e:
        print(f"❌ /cancel-reminder error: {e}")
        return jsonify({"error": str(e)}), 500

# ═══════════════════════════════════════════════════════════════
#  ERROR HANDLERS
# ═══════════════════════════════════════════════════════════════
@app.errorhandler(404)
def not_found(e):
    return jsonify({"error": "Endpoint not found"}), 404

@app.errorhandler(500)
def server_error(e):
    return jsonify({"error": "Internal server error"}), 500

@app.errorhandler(413)
def too_large(e):
    return jsonify({"error": "File too large. Maximum size is 20MB."}), 413

# ═══════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════
if __name__ == '__main__':
    # For local development only
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)